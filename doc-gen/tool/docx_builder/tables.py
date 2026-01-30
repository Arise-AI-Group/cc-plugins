"""Table support for DOCX builder with professional styling."""

from pathlib import Path
from typing import Any

from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from .core import load_document, save_document
from ..builder_common.presets import get_color, get_typography


def _parse_color(color_str: str) -> RGBColor:
    """Parse hex color string to RGBColor."""
    color_str = color_str.lstrip("#")
    return RGBColor(
        int(color_str[0:2], 16),
        int(color_str[2:4], 16),
        int(color_str[4:6], 16),
    )


def _hex_to_xml_color(color_str: str) -> str:
    """Convert hex color to XML format (without #)."""
    return color_str.lstrip("#").upper()


def _set_cell_shading(cell, color: str):
    """Set cell background color using XML."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{_hex_to_xml_color(color)}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _set_cell_borders(cell, color: str = "#CCCCCC", width: int = 1):
    """Set cell borders using XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Create borders element
    tcBorders = parse_xml(
        f'''<w:tcBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="{width * 8}" w:color="{_hex_to_xml_color(color)}"/>
            <w:left w:val="single" w:sz="{width * 8}" w:color="{_hex_to_xml_color(color)}"/>
            <w:bottom w:val="single" w:sz="{width * 8}" w:color="{_hex_to_xml_color(color)}"/>
            <w:right w:val="single" w:sz="{width * 8}" w:color="{_hex_to_xml_color(color)}"/>
        </w:tcBorders>'''
    )
    tcPr.append(tcBorders)


def _set_cell_width(cell, width_inches: float):
    """Set cell width."""
    cell.width = Inches(width_inches)


def add_table(
    doc_path: str | Path,
    data: list[list[str]],
    header_row: bool = True,
    column_widths: list[float] | None = None,
    style: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """Add a table to the document with professional styling.

    Args:
        doc_path: Path to the DOCX file
        data: 2D array of cell content [[row1], [row2], ...]
        header_row: If True, first row is styled as header
        column_widths: Optional list of column widths in inches
        style: Optional style dict with:
            - header_bg: Header background color (default: "#1E3A5F")
            - header_color: Header text color (default: "#FFFFFF")
            - border_color: Border color (default: "#CCCCCC")
            - alternating_rows: Enable alternating row colors (default: False)
            - alternating_color: Color for alternating rows (default: "#F5F5F5")
            - total_row: If True, style last row as total (default: False)

    Returns:
        Dict with status and table info
    """
    doc, metadata = load_document(doc_path)
    preset = metadata.get("preset_data", {})

    if not data:
        return {"success": False, "error": "No data provided"}

    # Default style
    default_style = {
        "header_bg": get_color(preset, "primary"),
        "header_color": "#FFFFFF",
        "border_color": "#CCCCCC",
        "alternating_rows": False,
        "alternating_color": "#F5F5F5",
        "total_row": False,
    }
    style = {**default_style, **(style or {})}

    # Create table
    num_rows = len(data)
    num_cols = len(data[0]) if data else 0
    table = doc.add_table(rows=num_rows, cols=num_cols)

    # Set table width to full page width
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Get font settings from preset
    font_name = get_typography(preset, "font_family")
    body_size = get_typography(preset, "body_size")

    # Calculate column widths
    if column_widths:
        widths = column_widths
    else:
        # Default: equal widths totaling ~6.5 inches (letter page with 1" margins)
        total_width = 6.5
        widths = [total_width / num_cols] * num_cols

    # Fill table
    for row_idx, row_data in enumerate(data):
        row = table.rows[row_idx]

        # Determine row styling
        is_header = header_row and row_idx == 0
        is_total = style["total_row"] and row_idx == num_rows - 1
        is_alternating = style["alternating_rows"] and not is_header and not is_total and row_idx % 2 == 0

        for col_idx, cell_text in enumerate(row_data):
            if col_idx >= num_cols:
                break

            cell = row.cells[col_idx]

            # Set cell width
            if col_idx < len(widths):
                _set_cell_width(cell, widths[col_idx])

            # Clear default paragraph and add text
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(str(cell_text))

            # Set font
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
            run.font.size = Pt(body_size / 2)

            # Style based on row type
            if is_header or is_total:
                run.bold = True
                run.font.color.rgb = _parse_color(style["header_color"])
                _set_cell_shading(cell, style["header_bg"])
            elif is_alternating:
                _set_cell_shading(cell, style["alternating_color"])
                run.font.color.rgb = _parse_color(get_color(preset, "text"))
            else:
                run.font.color.rgb = _parse_color(get_color(preset, "text"))

            # Set borders
            _set_cell_borders(cell, style["border_color"])

            # Set cell padding
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = parse_xml(
                f'''<w:tcMar {nsdecls("w")}>
                    <w:top w:w="80" w:type="dxa"/>
                    <w:left w:w="120" w:type="dxa"/>
                    <w:bottom w:w="80" w:type="dxa"/>
                    <w:right w:w="120" w:type="dxa"/>
                </w:tcMar>'''
            )
            tcPr.append(tcMar)

    save_document(doc, doc_path)

    return {
        "success": True,
        "element": "table",
        "rows": num_rows,
        "columns": num_cols,
        "header_row": header_row,
    }


def add_simple_table(
    doc_path: str | Path,
    data: list[list[str]],
    borders: bool = True,
) -> dict[str, Any]:
    """Add a simple table without special styling.

    Args:
        doc_path: Path to the DOCX file
        data: 2D array of cell content
        borders: Whether to show borders

    Returns:
        Dict with status
    """
    doc, metadata = load_document(doc_path)
    preset = metadata.get("preset_data", {})

    if not data:
        return {"success": False, "error": "No data provided"}

    num_rows = len(data)
    num_cols = len(data[0]) if data else 0
    table = doc.add_table(rows=num_rows, cols=num_cols)

    font_name = get_typography(preset, "font_family")
    body_size = get_typography(preset, "body_size")

    for row_idx, row_data in enumerate(data):
        row = table.rows[row_idx]
        for col_idx, cell_text in enumerate(row_data):
            if col_idx >= num_cols:
                break
            cell = row.cells[col_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(str(cell_text))
            run.font.name = font_name
            run.font.size = Pt(body_size / 2)

            if borders:
                _set_cell_borders(cell, "#CCCCCC")

    save_document(doc, doc_path)

    return {"success": True, "element": "simple_table", "rows": num_rows, "columns": num_cols}

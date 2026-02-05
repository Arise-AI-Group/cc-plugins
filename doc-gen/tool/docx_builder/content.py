"""Content elements for DOCX builder: headings, paragraphs, and lists."""

import re
from pathlib import Path
from typing import Any

from docx.shared import Pt

from .core import load_document, save_document
from ..styling import StyleResolver, apply_run_style, get_alignment


def _parse_markdown_text(text: str) -> list[tuple[str, dict]]:
    """Parse simple markdown formatting (bold, italic) into segments.

    Args:
        text: Text with optional **bold** and *italic* markers

    Returns:
        List of (text, {bold: bool, italic: bool}) tuples
    """
    segments = []
    # Pattern to match **bold**, *italic*, or plain text
    pattern = r"(\*\*[^*]+\*\*|\*[^*]+\*|[^*]+)"

    for match in re.finditer(pattern, text):
        segment = match.group(0)
        if segment.startswith("**") and segment.endswith("**"):
            segments.append((segment[2:-2], {"bold": True, "italic": False}))
        elif segment.startswith("*") and segment.endswith("*"):
            segments.append((segment[1:-1], {"bold": False, "italic": True}))
        else:
            segments.append((segment, {"bold": False, "italic": False}))

    return segments


def add_heading(
    doc_path: str | Path,
    text: str,
    level: int = 1,
    color: str | None = None,
    font: str | None = None,
    alignment: str = "left",
    spacing_before: float | None = None,
    spacing_after: float | None = None,
) -> dict[str, Any]:
    """Add a heading to the document.

    Args:
        doc_path: Path to the DOCX file
        text: Heading text
        level: Heading level (1-3)
        color: Optional hex color (e.g., "#1E3A5F") or named color ("primary")
        font: Optional font name
        alignment: "left", "center", "right", or "justify"
        spacing_before: Optional spacing before in points
        spacing_after: Optional spacing after in points

    Returns:
        Dict with status
    """
    doc, metadata = load_document(doc_path)
    style_config = metadata.get("style_config", metadata.get("preset_data", {}))
    resolver = StyleResolver(style_config)

    # Determine element name for style lookup
    element_name = f"heading{level}" if level > 1 else "title"
    element_style = resolver.get_element_style(element_name)
    font_config = element_style.get("font", {}).copy()

    # Apply overrides
    if font:
        font_config["name"] = font
    if color:
        font_config["color"] = color

    # Create paragraph and add text
    para = doc.add_paragraph()

    # Parse text for markdown formatting
    segments = _parse_markdown_text(text)
    for seg_text, formatting in segments:
        run = para.add_run(seg_text)

        # Build font config for this run
        run_font_config = font_config.copy()
        # Headings are always bold (unless formatting overrides)
        run_font_config["bold"] = run_font_config.get("bold", True) or formatting.get("bold", False)
        run_font_config["italic"] = formatting.get("italic", False)

        apply_run_style(run, run_font_config, resolver)

    # Set alignment
    style_alignment = element_style.get("alignment", alignment)
    para.alignment = get_alignment(style_alignment)

    # Set spacing
    pf = para.paragraph_format
    if spacing_before is not None:
        pf.space_before = Pt(spacing_before)
    else:
        pf.space_before = resolver.get_spacing(element_name, "before")

    if spacing_after is not None:
        pf.space_after = Pt(spacing_after)
    else:
        pf.space_after = resolver.get_spacing(element_name, "after")

    save_document(doc, doc_path)

    return {"success": True, "element": "heading", "level": level, "text": text[:50]}


def add_paragraph(
    doc_path: str | Path,
    text: str,
    bold: bool = False,
    italic: bool = False,
    color: str | None = None,
    font: str | None = None,
    font_size: float | None = None,
    alignment: str = "left",
    spacing_after: float | None = None,
) -> dict[str, Any]:
    """Add a paragraph to the document.

    Args:
        doc_path: Path to the DOCX file
        text: Paragraph text (supports **bold** and *italic* markdown)
        bold: Make entire paragraph bold
        italic: Make entire paragraph italic
        color: Optional hex color or named color
        font: Optional font name
        font_size: Optional font size in points
        alignment: "left", "center", "right", or "justify"
        spacing_after: Optional spacing after in points

    Returns:
        Dict with status
    """
    doc, metadata = load_document(doc_path)
    style_config = metadata.get("style_config", metadata.get("preset_data", {}))
    resolver = StyleResolver(style_config)

    # Get body style
    element_style = resolver.get_element_style("body")
    font_config = element_style.get("font", {}).copy()

    # Apply overrides
    if font:
        font_config["name"] = font
    if font_size is not None:
        font_config["size"] = font_size
    if color:
        font_config["color"] = color

    para = doc.add_paragraph()

    # Parse text for markdown formatting
    segments = _parse_markdown_text(text)
    for seg_text, formatting in segments:
        run = para.add_run(seg_text)

        # Build font config for this run
        run_font_config = font_config.copy()
        run_font_config["bold"] = bold or formatting.get("bold", False)
        run_font_config["italic"] = italic or formatting.get("italic", False)

        apply_run_style(run, run_font_config, resolver)

    # Set alignment
    para.alignment = get_alignment(alignment)

    # Set spacing
    pf = para.paragraph_format
    if spacing_after is not None:
        pf.space_after = Pt(spacing_after)
    else:
        pf.space_after = resolver.get_spacing("body", "after")

    save_document(doc, doc_path)

    return {"success": True, "element": "paragraph", "text": text[:50]}


def add_bullet_list(
    doc_path: str | Path,
    items: list[str],
    style: str = "bullet",
    color: str | None = None,
    font: str | None = None,
    font_size: float | None = None,
) -> dict[str, Any]:
    """Add a bullet list to the document.

    Args:
        doc_path: Path to the DOCX file
        items: List of items (each can contain **bold** markdown)
        style: "bullet", "check", "cross", or "dash"
        color: Optional hex color or named color
        font: Optional font name
        font_size: Optional font size in points

    Returns:
        Dict with status
    """
    doc, metadata = load_document(doc_path)
    style_config = metadata.get("style_config", metadata.get("preset_data", {}))
    resolver = StyleResolver(style_config)

    # Get bullet style and list config
    element_style = resolver.get_element_style("bullet")
    font_config = element_style.get("font", {}).copy()
    list_config = resolver.get_list_config(style)

    # Get bullet symbol and color
    bullet_char = list_config.get("symbol", "\u2022")
    bullet_color = list_config.get("color", "text")

    # Apply overrides
    if font:
        font_config["name"] = font
    if font_size is not None:
        font_config["size"] = font_size
    if color:
        font_config["color"] = color

    for item_text in items:
        para = doc.add_paragraph()

        # Add bullet character with its color
        bullet_run = para.add_run(f"{bullet_char} ")
        bullet_font_config = font_config.copy()
        bullet_font_config["color"] = bullet_color
        bullet_font_config["bold"] = True
        apply_run_style(bullet_run, bullet_font_config, resolver)

        # Parse and add item text
        segments = _parse_markdown_text(item_text)
        for seg_text, formatting in segments:
            run = para.add_run(seg_text)

            run_font_config = font_config.copy()
            run_font_config["bold"] = formatting.get("bold", False)
            run_font_config["italic"] = formatting.get("italic", False)

            apply_run_style(run, run_font_config, resolver)

        # Set indentation for list appearance
        pf = para.paragraph_format
        pf.left_indent = Pt(20)
        pf.space_after = resolver.get_spacing("bullet", "after", 4)

    save_document(doc, doc_path)

    return {"success": True, "element": "bullet_list", "count": len(items)}


def add_numbered_list(
    doc_path: str | Path,
    items: list[str],
    start: int = 1,
    color: str | None = None,
    font: str | None = None,
    font_size: float | None = None,
) -> dict[str, Any]:
    """Add a numbered list to the document.

    Args:
        doc_path: Path to the DOCX file
        items: List of items (each can contain **bold** markdown)
        start: Starting number
        color: Optional hex color or named color
        font: Optional font name
        font_size: Optional font size in points

    Returns:
        Dict with status
    """
    doc, metadata = load_document(doc_path)
    style_config = metadata.get("style_config", metadata.get("preset_data", {}))
    resolver = StyleResolver(style_config)

    # Get bullet style (used for numbered lists too)
    element_style = resolver.get_element_style("bullet")
    font_config = element_style.get("font", {}).copy()

    # Apply overrides
    if font:
        font_config["name"] = font
    if font_size is not None:
        font_config["size"] = font_size
    if color:
        font_config["color"] = color

    for i, item_text in enumerate(items, start=start):
        para = doc.add_paragraph()

        # Add number
        num_run = para.add_run(f"{i}. ")
        num_font_config = font_config.copy()
        num_font_config["bold"] = True
        apply_run_style(num_run, num_font_config, resolver)

        # Parse and add item text
        segments = _parse_markdown_text(item_text)
        for seg_text, formatting in segments:
            run = para.add_run(seg_text)

            run_font_config = font_config.copy()
            run_font_config["bold"] = formatting.get("bold", False)
            run_font_config["italic"] = formatting.get("italic", False)

            apply_run_style(run, run_font_config, resolver)

        # Set indentation
        pf = para.paragraph_format
        pf.left_indent = Pt(20)
        pf.space_after = resolver.get_spacing("bullet", "after", 4)

    save_document(doc, doc_path)

    return {"success": True, "element": "numbered_list", "count": len(items)}


def _render_template_vars(text: str, context: dict[str, Any]) -> str:
    """Render Jinja2-style template variables in text.

    Supports:
    - {{ variable }} - simple variable
    - {{ colors.primary }} - nested access

    Args:
        text: Text with template variables
        context: Dict of variable values

    Returns:
        Rendered text
    """
    import re

    def replace_var(match: re.Match) -> str:
        var_path = match.group(1).strip()
        parts = var_path.split(".")

        value = context
        for part in parts:
            if isinstance(value, dict):
                value = value.get(part, "")
            else:
                return match.group(0)  # Keep original if not found

        return str(value) if value else ""

    return re.sub(r"\{\{\s*([^}]+)\s*\}\}", replace_var, text)


def setup_header_footer(
    doc_path: str | Path,
    header_config: dict[str, Any] | None = None,
    footer_config: dict[str, Any] | None = None,
    context: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """Set up header and/or footer for the document.

    Args:
        doc_path: Path to the DOCX file
        header_config: Header configuration dict with:
            - enabled: bool
            - text: str (supports {{ variables }})
            - position: "left", "center", "right"
            - font_size: float (points)
            - color: hex or named color
        footer_config: Footer configuration dict (same fields as header)
        context: Variables for template rendering (e.g., {"page": "{{ page }}"})

    Returns:
        Dict with status
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc, metadata = load_document(doc_path)
    style_config = metadata.get("style_config", metadata.get("preset_data", {}))
    resolver = StyleResolver(style_config)

    # Build context from style config
    render_context = {
        "company": style_config.get("contact", {}).get("company", ""),
        "company_short": style_config.get("contact", {}).get("company_short", ""),
        "tagline": style_config.get("contact", {}).get("tagline", ""),
        "email": style_config.get("contact", {}).get("email", ""),
        "phone": style_config.get("contact", {}).get("phone", ""),
        "website": style_config.get("contact", {}).get("website", ""),
        "colors": style_config.get("colors", {}),
    }
    if context:
        render_context.update(context)

    section = doc.sections[0]

    results = {"success": True, "header": None, "footer": None}

    # Set up header
    if header_config and header_config.get("enabled", False):
        header = section.header
        header.is_linked_to_previous = False

        # Clear existing content
        for para in header.paragraphs:
            p = para._element
            p.getparent().remove(p)

        # Add header paragraph
        para = header.add_paragraph()
        text = header_config.get("text", "")
        rendered_text = _render_template_vars(text, render_context)

        run = para.add_run(rendered_text)

        # Style the run
        font_config = {
            "size": header_config.get("font_size", 10),
            "color": header_config.get("color", "muted"),
        }
        apply_run_style(run, font_config, resolver)

        # Set alignment
        position = header_config.get("position", "right")
        if position == "left":
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif position == "center":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif position == "right":
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        results["header"] = {"enabled": True, "text": rendered_text}

    # Set up footer
    if footer_config and footer_config.get("enabled", True):
        footer = section.footer
        footer.is_linked_to_previous = False

        # Clear existing content
        for para in footer.paragraphs:
            p = para._element
            p.getparent().remove(p)

        # Add footer paragraph
        para = footer.add_paragraph()
        text = footer_config.get("text", "Page {{ page }}")

        # Handle {{ page }} specially - insert page number field
        if "{{ page }}" in text:
            parts = text.split("{{ page }}")
            for i, part in enumerate(parts):
                if part:
                    rendered_part = _render_template_vars(part, render_context)
                    run = para.add_run(rendered_part)
                    font_config = {
                        "size": footer_config.get("font_size", 9),
                        "color": footer_config.get("color", "muted"),
                    }
                    apply_run_style(run, font_config, resolver)

                # Add page number field (except after last part)
                if i < len(parts) - 1:
                    _add_page_number_field(para, resolver, footer_config)
        else:
            rendered_text = _render_template_vars(text, render_context)
            run = para.add_run(rendered_text)
            font_config = {
                "size": footer_config.get("font_size", 9),
                "color": footer_config.get("color", "muted"),
            }
            apply_run_style(run, font_config, resolver)

        # Set alignment
        position = footer_config.get("position", "center")
        if position == "left":
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif position == "center":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif position == "right":
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        results["footer"] = {"enabled": True, "text": text}

    save_document(doc, doc_path)

    return results


def _add_page_number_field(para, resolver: StyleResolver, config: dict[str, Any]) -> None:
    """Add a PAGE field to a paragraph.

    Args:
        para: docx paragraph object
        resolver: StyleResolver for styling
        config: Footer/header config for font settings
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Create a run for the field
    run = para.add_run()

    # Build the field elements
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')

    instr_text = OxmlElement('w:instrText')
    instr_text.text = "PAGE"
    instr_text.set(qn('xml:space'), 'preserve')

    fld_char_separate = OxmlElement('w:fldChar')
    fld_char_separate.set(qn('w:fldCharType'), 'separate')

    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')

    # Add to run
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)

    # Style the run
    font_config = {
        "size": config.get("font_size", 9),
        "color": config.get("color", "muted"),
    }
    apply_run_style(run, font_config, resolver)

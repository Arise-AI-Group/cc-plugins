"""Core document lifecycle operations for DOCX builder.

Handles document creation, loading, saving, and metadata management.
"""

import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor

from ..styling import load_style, StyleResolver, DEFAULT_STYLE


def _get_metadata_path(doc_path: Path) -> Path:
    """Get the sidecar metadata file path for a document."""
    return doc_path.with_suffix(".docx.meta.json")


def _parse_color(color_str: str) -> RGBColor:
    """Parse hex color string to RGBColor.

    Args:
        color_str: Hex color like "#1E3A5F" or "1E3A5F"

    Returns:
        RGBColor object
    """
    color_str = color_str.lstrip("#")
    return RGBColor(
        int(color_str[0:2], 16),
        int(color_str[2:4], 16),
        int(color_str[4:6], 16),
    )


def create_document(
    output_path: str | Path,
    preset: str | None = None,
    page_size: str = "letter",
    margins: dict[str, float] | None = None,
) -> dict[str, Any]:
    """Create a new DOCX document with optional preset.

    Args:
        output_path: Path for the new document
        preset: Optional preset/style name (e.g., "40hero", "professional", "default")
        page_size: Page size ("letter" or "a4")
        margins: Optional margins dict {top, right, bottom, left} in inches

    Returns:
        Dict with status and document info
    """
    output_path = Path(output_path)

    # Load style configuration (supports both JSON styles and legacy YAML presets)
    style_config = load_style(preset)
    resolver = StyleResolver(style_config)

    # Create document
    doc = Document()

    # Set page size
    section = doc.sections[0]
    page_config = style_config.get("page", {})
    config_page_size = page_config.get("size", page_size)

    if config_page_size == "letter":
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    elif config_page_size == "a4":
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)

    # Set margins - prefer explicit parameter, then style config, then defaults
    if margins:
        section.top_margin = Inches(margins.get("top", 1.0))
        section.right_margin = Inches(margins.get("right", 1.0))
        section.bottom_margin = Inches(margins.get("bottom", 1.0))
        section.left_margin = Inches(margins.get("left", 1.0))
    else:
        # Use margins from style config
        style_margins = resolver.get_page_margins()
        section.top_margin = style_margins.get("top", Inches(1.0))
        section.right_margin = style_margins.get("right", Inches(1.0))
        section.bottom_margin = style_margins.get("bottom", Inches(1.0))
        section.left_margin = style_margins.get("left", Inches(1.0))

    # Set default font based on style
    body_font = resolver.get_body_font()
    style = doc.styles["Normal"]
    style.font.name = body_font.get("name", "Arial")
    style.font.size = Pt(body_font.get("size", 11))

    # Save document
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    # Save metadata sidecar with style configuration
    metadata = {
        "preset": preset or "default",
        "style_config": style_config,
        # Keep legacy key for backward compatibility
        "preset_data": style_config,
        "page_size": config_page_size,
        "margins": margins or {
            "top": float(str(style_margins.get("top", Inches(1.0))).replace(" inches", "")) if not margins else margins.get("top", 1.0),
            "right": 1.0,
            "bottom": 1.0,
            "left": 1.0,
        },
    }
    metadata_path = _get_metadata_path(output_path)
    with open(metadata_path, "w") as f:
        json.dump(metadata, f, indent=2)

    return {
        "success": True,
        "path": str(output_path),
        "preset": preset or "default",
        "page_size": config_page_size,
    }


def load_document(doc_path: str | Path) -> tuple[Document, dict[str, Any]]:
    """Load a document and its metadata.

    Args:
        doc_path: Path to the DOCX file

    Returns:
        Tuple of (Document object, metadata dict)

    Raises:
        FileNotFoundError: If document doesn't exist
    """
    doc_path = Path(doc_path)
    if not doc_path.exists():
        raise FileNotFoundError(f"Document not found: {doc_path}")

    doc = Document(str(doc_path))

    # Load metadata
    metadata_path = _get_metadata_path(doc_path)
    if metadata_path.exists():
        with open(metadata_path) as f:
            metadata = json.load(f)
        # Ensure style_config is available (support both old and new formats)
        if "style_config" not in metadata and "preset_data" in metadata:
            metadata["style_config"] = metadata["preset_data"]
    else:
        # No metadata, use defaults
        metadata = {
            "preset": "default",
            "style_config": DEFAULT_STYLE,
            "preset_data": DEFAULT_STYLE,  # Legacy compatibility
            "page_size": "letter",
            "margins": {"top": 1.0, "right": 1.0, "bottom": 1.0, "left": 1.0},
        }

    return doc, metadata


def save_document(doc: Document, doc_path: str | Path) -> dict[str, Any]:
    """Save a document.

    Args:
        doc: Document object to save
        doc_path: Path to save to

    Returns:
        Dict with status
    """
    doc_path = Path(doc_path)
    doc.save(str(doc_path))
    return {"success": True, "path": str(doc_path)}


def get_document_metadata(doc_path: str | Path) -> dict[str, Any]:
    """Get metadata for a document.

    Args:
        doc_path: Path to the DOCX file

    Returns:
        Metadata dict
    """
    doc_path = Path(doc_path)
    metadata_path = _get_metadata_path(doc_path)

    if metadata_path.exists():
        with open(metadata_path) as f:
            return json.load(f)

    return {
        "preset": "default",
        "page_size": "letter",
        "note": "No metadata file found, using defaults",
    }


def finalize_document(doc_path: str | Path, cleanup_metadata: bool = False) -> dict[str, Any]:
    """Finalize a document (optional validation/cleanup).

    Args:
        doc_path: Path to the DOCX file
        cleanup_metadata: If True, remove the sidecar metadata file

    Returns:
        Dict with status
    """
    doc_path = Path(doc_path)
    if not doc_path.exists():
        return {"success": False, "error": f"Document not found: {doc_path}"}

    # Load and re-save to ensure validity
    doc = Document(str(doc_path))
    doc.save(str(doc_path))

    # Optionally clean up metadata
    if cleanup_metadata:
        metadata_path = _get_metadata_path(doc_path)
        if metadata_path.exists():
            metadata_path.unlink()

    return {"success": True, "path": str(doc_path), "finalized": True}

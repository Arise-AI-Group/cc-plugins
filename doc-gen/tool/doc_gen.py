#!/usr/bin/env python3
"""Document generation CLI using WeasyPrint and pure Python libraries.

Converts HTML/Markdown to PDF and DOCX with Jinja2 templating support.

CLI Usage:
    # HTML to PDF (WeasyPrint)
    ./run tool/doc_gen.py pdf input.html -o output.pdf
    ./run tool/doc_gen.py pdf input.html --style quote -o output.pdf
    ./run tool/doc_gen.py pdf input.html --css custom.css -o output.pdf

    # With Jinja2 template variables
    ./run tool/doc_gen.py pdf template.html -o output.pdf --var client="Acme Corp"
    ./run tool/doc_gen.py pdf template.html -o output.pdf --vars data.json

    # Markdown to DOCX
    ./run tool/doc_gen.py docx input.md -o output.docx

    # Markdown to PDF
    ./run tool/doc_gen.py pdf input.md -o output.pdf

Module Usage:
    from tool.doc_gen import DocGenerator

    gen = DocGenerator()
    gen.html_to_pdf("input.html", "output.pdf", style="quote")
    gen.markdown_to_docx("input.md", "output.docx")
"""

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any

from jinja2 import Environment, FileSystemLoader, BaseLoader


class DocGenError(Exception):
    """Base exception for document generation errors."""
    pass


class DependencyError(DocGenError):
    """Raised when a required dependency is missing."""
    pass


class TemplateError(DocGenError):
    """Raised when template processing fails."""
    pass


class ConversionError(DocGenError):
    """Raised when document conversion fails."""
    pass


def get_plugin_dir() -> Path:
    """Get the plugin directory (parent of tool/)."""
    return Path(__file__).parent.parent


def get_styles_dir() -> Path:
    """Get the styles directory."""
    return get_plugin_dir() / "styles"


def get_templates_dir() -> Path:
    """Get the templates directory."""
    return get_plugin_dir() / "templates"


class DocGenerator:
    """Document generator using WeasyPrint and pure Python libraries."""

    def __init__(self):
        self.styles_dir = get_styles_dir()
        self.templates_dir = get_templates_dir()

    def _check_weasyprint(self) -> None:
        """Check if WeasyPrint is available."""
        try:
            import weasyprint  # noqa: F401
        except ImportError:
            raise DependencyError(
                "WeasyPrint not installed. Run: pip install weasyprint"
            )

    def _render_template(
        self,
        content: str,
        variables: dict[str, Any],
        base_path: Path | None = None
    ) -> str:
        """Render Jinja2 template with variables.

        Args:
            content: Template content string
            variables: Dictionary of template variables
            base_path: Base path for template includes (optional)

        Returns:
            Rendered template string
        """
        if not variables:
            return content

        try:
            if base_path:
                env = Environment(loader=FileSystemLoader(str(base_path)))
                template = env.from_string(content)
            else:
                env = Environment(loader=BaseLoader())
                template = env.from_string(content)

            return template.render(**variables)
        except Exception as e:
            raise TemplateError(f"Template rendering failed: {e}")

    def _get_style_css(self, style_name: str) -> str | None:
        """Get CSS content for a built-in style.

        Args:
            style_name: Name of the style (e.g., 'quote', 'report', 'invoice')

        Returns:
            CSS content string or None if style not found
        """
        style_file = self.styles_dir / f"{style_name}.css"
        if style_file.exists():
            return style_file.read_text()
        return None

    def html_to_pdf(
        self,
        input_path: str | Path,
        output_path: str | Path,
        css_path: str | Path | None = None,
        style: str | None = None,
        variables: dict[str, Any] | None = None
    ) -> Path:
        """Convert HTML to PDF using WeasyPrint.

        Args:
            input_path: Path to input HTML file
            output_path: Path for output PDF file
            css_path: Optional path to custom CSS file
            style: Optional built-in style name (quote, report, invoice)
            variables: Optional Jinja2 template variables

        Returns:
            Path to the created PDF file
        """
        self._check_weasyprint()
        from weasyprint import HTML, CSS

        input_path = Path(input_path)
        output_path = Path(output_path)

        if not input_path.exists():
            raise ConversionError(f"Input file not found: {input_path}")

        # Read and optionally render template
        content = input_path.read_text()
        if variables:
            content = self._render_template(
                content, variables, input_path.parent
            )

        # Build stylesheets list
        stylesheets = []

        # Add built-in style if specified
        if style:
            style_css = self._get_style_css(style)
            if style_css:
                stylesheets.append(CSS(string=style_css))
            else:
                print(f"Warning: Style '{style}' not found, skipping", file=sys.stderr)

        # Add custom CSS if specified
        if css_path:
            css_path = Path(css_path)
            if css_path.exists():
                stylesheets.append(CSS(filename=str(css_path)))
            else:
                raise ConversionError(f"CSS file not found: {css_path}")

        # Create PDF
        html = HTML(string=content, base_url=str(input_path.parent))
        html.write_pdf(
            str(output_path),
            stylesheets=stylesheets if stylesheets else None
        )

        return output_path

    def markdown_to_pdf(
        self,
        input_path: str | Path,
        output_path: str | Path,
        variables: dict[str, Any] | None = None,
        style: str | None = None
    ) -> Path:
        """Convert Markdown to PDF using mistune + WeasyPrint.

        Args:
            input_path: Path to input Markdown file
            output_path: Path for output PDF file
            variables: Optional Jinja2 template variables
            style: Optional built-in style name

        Returns:
            Path to the created PDF file
        """
        self._check_weasyprint()
        import mistune
        from weasyprint import HTML, CSS

        input_path = Path(input_path)
        output_path = Path(output_path)

        if not input_path.exists():
            raise ConversionError(f"Input file not found: {input_path}")

        # Read and optionally render template
        content = input_path.read_text()
        if variables:
            content = self._render_template(
                content, variables, input_path.parent
            )

        # Convert Markdown to HTML
        html_content = mistune.html(content)

        # Wrap in full HTML structure
        full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
            font-size: 11pt;
            line-height: 1.6;
            max-width: 7in;
            margin: 0.75in auto;
            color: #333;
        }}
        h1 {{ font-size: 24pt; margin-top: 0; margin-bottom: 16pt; font-weight: 600; }}
        h2 {{ font-size: 18pt; margin-top: 24pt; margin-bottom: 12pt; font-weight: 600; }}
        h3 {{ font-size: 14pt; margin-top: 20pt; margin-bottom: 10pt; font-weight: 600; }}
        h4, h5, h6 {{ font-size: 12pt; margin-top: 16pt; margin-bottom: 8pt; font-weight: 600; }}
        p {{ margin: 0 0 12pt 0; }}
        pre {{ background: #f5f5f5; padding: 12pt; border-radius: 4pt; overflow-x: auto; }}
        code {{ font-family: "SF Mono", Monaco, monospace; font-size: 10pt; }}
        pre code {{ background: none; }}
        code:not(pre code) {{ background: #f0f0f0; padding: 2pt 4pt; border-radius: 3pt; }}
        blockquote {{ margin: 12pt 0; padding-left: 16pt; border-left: 3pt solid #ddd; color: #666; }}
        table {{ border-collapse: collapse; width: 100%; margin: 16pt 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8pt 12pt; text-align: left; }}
        th {{ background: #f5f5f5; font-weight: 600; }}
        ul, ol {{ margin: 0 0 12pt 0; padding-left: 24pt; }}
        li {{ margin-bottom: 4pt; }}
        hr {{ border: none; border-top: 1px solid #ddd; margin: 24pt 0; }}
        a {{ color: #0066cc; text-decoration: none; }}
        img {{ max-width: 100%; height: auto; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""

        # Build stylesheets list
        stylesheets = []
        if style:
            style_css = self._get_style_css(style)
            if style_css:
                stylesheets.append(CSS(string=style_css))

        # Convert to PDF
        HTML(string=full_html, base_url=str(input_path.parent)).write_pdf(
            str(output_path),
            stylesheets=stylesheets if stylesheets else None
        )

        return output_path

    def markdown_to_docx(
        self,
        input_path: str | Path,
        output_path: str | Path,
        variables: dict[str, Any] | None = None
    ) -> Path:
        """Convert Markdown to DOCX using mistune + python-docx.

        Args:
            input_path: Path to input Markdown file
            output_path: Path for output DOCX file
            variables: Optional Jinja2 template variables

        Returns:
            Path to the created DOCX file
        """
        import mistune
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        input_path = Path(input_path)
        output_path = Path(output_path)

        if not input_path.exists():
            raise ConversionError(f"Input file not found: {input_path}")

        # Read and optionally render template
        content = input_path.read_text()
        if variables:
            content = self._render_template(
                content, variables, input_path.parent
            )

        # Parse markdown into AST
        md = mistune.create_markdown(renderer=None)
        tokens = md(content)

        # Create DOCX document
        doc = Document()

        # Process tokens
        self._process_markdown_tokens(doc, tokens)

        # Save document
        doc.save(str(output_path))

        return output_path

    def _process_markdown_tokens(self, doc, tokens, list_level=0):
        """Process markdown tokens and add to document.

        Args:
            doc: python-docx Document
            tokens: List of mistune tokens
            list_level: Current list nesting level
        """
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        for token in tokens:
            token_type = token.get('type')

            if token_type == 'heading':
                level = token.get('attrs', {}).get('level', 1)
                text = self._extract_text_from_children(token.get('children', []))
                heading = doc.add_heading(text, level=level)

            elif token_type == 'paragraph':
                text = self._extract_text_from_children(token.get('children', []))
                if text.strip():
                    doc.add_paragraph(text)

            elif token_type == 'code_block':
                code = token.get('raw', '')
                para = doc.add_paragraph()
                run = para.add_run(code)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                para.paragraph_format.left_indent = Inches(0.25)

            elif token_type == 'block_quote':
                children = token.get('children', [])
                for child in children:
                    if child.get('type') == 'paragraph':
                        text = self._extract_text_from_children(child.get('children', []))
                        para = doc.add_paragraph(text)
                        para.paragraph_format.left_indent = Inches(0.5)
                        para.style = 'Quote' if 'Quote' in [s.name for s in doc.styles] else None

            elif token_type == 'list':
                ordered = token.get('attrs', {}).get('ordered', False)
                items = token.get('children', [])
                for i, item in enumerate(items):
                    if item.get('type') == 'list_item':
                        text = self._extract_text_from_children(
                            item.get('children', [{}])[0].get('children', [])
                            if item.get('children') and item['children'][0].get('type') == 'paragraph'
                            else []
                        )
                        if ordered:
                            para = doc.add_paragraph(f"{i+1}. {text}")
                        else:
                            para = doc.add_paragraph(text, style='List Bullet')

            elif token_type == 'thematic_break':
                para = doc.add_paragraph()
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after = Pt(12)
                # Add horizontal line
                run = para.add_run('_' * 50)
                run.font.color.rgb = None  # Use theme color

    def _extract_text_from_children(self, children):
        """Extract plain text from token children.

        Args:
            children: List of child tokens

        Returns:
            Extracted text string
        """
        text_parts = []
        for child in children:
            if child.get('type') == 'text':
                text_parts.append(child.get('raw', ''))
            elif child.get('type') == 'codespan':
                text_parts.append(child.get('raw', ''))
            elif child.get('type') in ('strong', 'emphasis', 'link'):
                text_parts.append(self._extract_text_from_children(child.get('children', [])))
            elif child.get('type') == 'softbreak':
                text_parts.append(' ')
        return ''.join(text_parts)

    def html_to_docx(
        self,
        input_path: str | Path,
        output_path: str | Path,
        variables: dict[str, Any] | None = None
    ) -> Path:
        """Convert HTML to DOCX using python-docx with HTML parsing.

        Args:
            input_path: Path to input HTML file
            output_path: Path for output DOCX file
            variables: Optional Jinja2 template variables

        Returns:
            Path to the created DOCX file
        """
        from docx import Document
        from docx.shared import Pt, Inches

        input_path = Path(input_path)
        output_path = Path(output_path)

        if not input_path.exists():
            raise ConversionError(f"Input file not found: {input_path}")

        # Read and optionally render template
        content = input_path.read_text()
        if variables:
            content = self._render_template(
                content, variables, input_path.parent
            )

        # Create document
        doc = Document()

        # Parse HTML and convert to DOCX
        self._parse_html_to_docx(doc, content)

        # Save document
        doc.save(str(output_path))

        return output_path

    def _parse_html_to_docx(self, doc, html_content):
        """Parse HTML content and add to document.

        Args:
            doc: python-docx Document
            html_content: HTML string
        """
        from docx.shared import Pt, Inches

        # Simple HTML to DOCX conversion using regex
        # Remove script and style tags
        html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)

        # Extract body content if present
        body_match = re.search(r'<body[^>]*>(.*?)</body>', html_content, flags=re.DOTALL | re.IGNORECASE)
        if body_match:
            html_content = body_match.group(1)

        # Process headings
        for level in range(1, 7):
            pattern = f'<h{level}[^>]*>(.*?)</h{level}>'
            html_content = re.sub(
                pattern,
                lambda m: f'\n[HEADING{level}]{self._strip_tags(m.group(1))}[/HEADING{level}]\n',
                html_content,
                flags=re.DOTALL | re.IGNORECASE
            )

        # Process paragraphs
        html_content = re.sub(
            r'<p[^>]*>(.*?)</p>',
            lambda m: f'\n[PARA]{self._strip_tags(m.group(1))}[/PARA]\n',
            html_content,
            flags=re.DOTALL | re.IGNORECASE
        )

        # Process list items
        html_content = re.sub(
            r'<li[^>]*>(.*?)</li>',
            lambda m: f'\n[LI]{self._strip_tags(m.group(1))}[/LI]\n',
            html_content,
            flags=re.DOTALL | re.IGNORECASE
        )

        # Process code blocks
        html_content = re.sub(
            r'<pre[^>]*>(.*?)</pre>',
            lambda m: f'\n[CODE]{self._strip_tags(m.group(1))}[/CODE]\n',
            html_content,
            flags=re.DOTALL | re.IGNORECASE
        )

        # Strip remaining tags and normalize whitespace
        remaining_text = self._strip_tags(html_content)

        # Process the marked content
        lines = html_content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Check for marked elements
            heading_match = re.match(r'\[HEADING(\d)\](.*?)\[/HEADING\d\]', line)
            if heading_match:
                level = int(heading_match.group(1))
                text = heading_match.group(2).strip()
                if text:
                    doc.add_heading(text, level=level)
                continue

            para_match = re.match(r'\[PARA\](.*?)\[/PARA\]', line)
            if para_match:
                text = para_match.group(1).strip()
                if text:
                    doc.add_paragraph(text)
                continue

            li_match = re.match(r'\[LI\](.*?)\[/LI\]', line)
            if li_match:
                text = li_match.group(1).strip()
                if text:
                    doc.add_paragraph(text, style='List Bullet')
                continue

            code_match = re.match(r'\[CODE\](.*?)\[/CODE\]', line, re.DOTALL)
            if code_match:
                text = code_match.group(1).strip()
                if text:
                    para = doc.add_paragraph()
                    run = para.add_run(text)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                continue

            # Plain text that's not inside markers
            text = self._strip_tags(line).strip()
            if text and not text.startswith('['):
                doc.add_paragraph(text)

    def _strip_tags(self, html):
        """Remove HTML tags from string.

        Args:
            html: HTML string

        Returns:
            Text with tags removed
        """
        # Decode HTML entities
        html = html.replace('&nbsp;', ' ')
        html = html.replace('&lt;', '<')
        html = html.replace('&gt;', '>')
        html = html.replace('&amp;', '&')
        html = html.replace('&quot;', '"')
        html = html.replace('&#39;', "'")

        # Remove tags
        clean = re.sub(r'<[^>]+>', '', html)

        # Normalize whitespace
        clean = re.sub(r'\s+', ' ', clean)

        return clean.strip()

    def list_styles(self) -> list[str]:
        """List available built-in styles.

        Returns:
            List of style names
        """
        if not self.styles_dir.exists():
            return []
        return [f.stem for f in self.styles_dir.glob("*.css")]


def parse_var(var_str: str) -> tuple[str, str]:
    """Parse a variable string like 'key=value'.

    Args:
        var_str: String in format 'key=value'

    Returns:
        Tuple of (key, value)
    """
    if '=' not in var_str:
        raise ValueError(f"Invalid variable format: {var_str} (expected key=value)")
    key, value = var_str.split('=', 1)
    return key.strip(), value.strip()


def build_parser() -> argparse.ArgumentParser:
    """Build the argument parser."""
    parser = argparse.ArgumentParser(
        prog="doc_gen",
        description="Document generation using WeasyPrint and pure Python libraries",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # HTML to PDF with WeasyPrint
  ./run tool/doc_gen.py pdf input.html -o output.pdf
  ./run tool/doc_gen.py pdf input.html --style quote -o output.pdf

  # With template variables
  ./run tool/doc_gen.py pdf template.html -o out.pdf --var client="Acme"
  ./run tool/doc_gen.py pdf template.html -o out.pdf --vars data.json

  # Markdown to DOCX
  ./run tool/doc_gen.py docx input.md -o output.docx

  # Markdown to PDF
  ./run tool/doc_gen.py pdf input.md -o output.pdf

  # List available styles
  ./run tool/doc_gen.py styles
"""
    )

    subparsers = parser.add_subparsers(dest="command", help="Available commands")

    # PDF command
    pdf_parser = subparsers.add_parser("pdf", help="Generate PDF from HTML or Markdown")
    pdf_parser.add_argument("input", help="Input file (HTML or Markdown)")
    pdf_parser.add_argument("-o", "--output", required=True, help="Output PDF file")
    pdf_parser.add_argument("--style", help="Built-in style (quote, report, invoice)")
    pdf_parser.add_argument("--css", help="Custom CSS file path")
    pdf_parser.add_argument(
        "--var",
        action="append",
        dest="vars",
        help="Template variable (key=value), can be repeated"
    )
    pdf_parser.add_argument(
        "--vars",
        dest="vars_file",
        help="JSON file with template variables"
    )

    # DOCX command
    docx_parser = subparsers.add_parser("docx", help="Generate DOCX from HTML or Markdown")
    docx_parser.add_argument("input", help="Input file (HTML or Markdown)")
    docx_parser.add_argument("-o", "--output", required=True, help="Output DOCX file")
    docx_parser.add_argument(
        "--var",
        action="append",
        dest="vars",
        help="Template variable (key=value), can be repeated"
    )
    docx_parser.add_argument(
        "--vars",
        dest="vars_file",
        help="JSON file with template variables"
    )

    # Styles command
    subparsers.add_parser("styles", help="List available built-in styles")

    return parser


def collect_variables(args) -> dict[str, Any]:
    """Collect template variables from CLI args and JSON file.

    Args:
        args: Parsed arguments with vars and vars_file attributes

    Returns:
        Dictionary of template variables
    """
    variables = {}

    # Load from JSON file first
    if hasattr(args, 'vars_file') and args.vars_file:
        vars_path = Path(args.vars_file)
        if not vars_path.exists():
            raise ValueError(f"Variables file not found: {vars_path}")
        with open(vars_path) as f:
            variables.update(json.load(f))

    # Override with CLI variables
    if hasattr(args, 'vars') and args.vars:
        for var_str in args.vars:
            key, value = parse_var(var_str)
            variables[key] = value

    return variables


def main():
    """Main CLI entry point."""
    parser = build_parser()
    args = parser.parse_args()

    if not args.command:
        parser.print_help()
        sys.exit(1)

    gen = DocGenerator()

    try:
        if args.command == "styles":
            styles = gen.list_styles()
            if styles:
                print("Available styles:")
                for style in styles:
                    print(f"  - {style}")
            else:
                print("No built-in styles found")
            return

        # Collect template variables
        variables = collect_variables(args)

        input_path = Path(args.input)
        output_path = Path(args.output)

        # Determine file type
        input_ext = input_path.suffix.lower()
        is_html = input_ext in ['.html', '.htm']
        is_markdown = input_ext in ['.md', '.markdown']

        if args.command == "pdf":
            if is_html:
                result = gen.html_to_pdf(
                    input_path,
                    output_path,
                    css_path=getattr(args, 'css', None),
                    style=getattr(args, 'style', None),
                    variables=variables
                )
            else:
                # Markdown to PDF
                result = gen.markdown_to_pdf(
                    input_path,
                    output_path,
                    variables=variables,
                    style=getattr(args, 'style', None)
                )

            print(json.dumps({
                "status": "success",
                "output": str(result),
                "format": "pdf"
            }, indent=2))

        elif args.command == "docx":
            if is_html:
                result = gen.html_to_docx(
                    input_path,
                    output_path,
                    variables=variables
                )
            else:
                result = gen.markdown_to_docx(
                    input_path,
                    output_path,
                    variables=variables
                )

            print(json.dumps({
                "status": "success",
                "output": str(result),
                "format": "docx"
            }, indent=2))

    except (DocGenError, ValueError) as e:
        print(json.dumps({
            "status": "error",
            "error": str(e)
        }, indent=2), file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()

"""Style resolver for resolving named colors and getting element styles.

The StyleResolver class provides methods to resolve named color references
(like "primary") to actual color values, and to retrieve style configurations
for different document elements.
"""

from typing import Any

from docx.shared import RGBColor, Pt, Inches


class StyleResolver:
    """Resolves style references to concrete values.

    Takes a unified style configuration and provides methods to resolve
    named color references, get element styles, and convert units.
    """

    def __init__(self, config: dict[str, Any]):
        """Initialize resolver with a style configuration.

        Args:
            config: Unified style configuration dict (from loader.load_style)
        """
        self.config = config
        self.colors = config.get("colors", {})
        self.fonts = config.get("fonts", {})
        self.styles = config.get("styles", {})
        self.tables = config.get("tables", {})
        self.lists = config.get("lists", {})
        self.page = config.get("page", {})
        self.contact = config.get("contact", {})

    def resolve_color(self, value: str | None, default: str = "#333333") -> RGBColor:
        """Resolve a color value to RGBColor.

        Accepts either:
        - Named color reference: "primary", "accent", "text", etc.
        - Hex color value: "#1E3A5F", "#333"

        Args:
            value: Color name or hex value
            default: Fallback hex color if value is None or not found

        Returns:
            RGBColor object
        """
        if not value:
            return self._hex_to_rgb(default)

        # If it's a hex color, use directly
        if value.startswith("#"):
            return self._hex_to_rgb(value)

        # Otherwise, look up in named colors
        resolved = self.colors.get(value)
        if resolved:
            # Recursively resolve in case of color aliasing
            if resolved.startswith("#"):
                return self._hex_to_rgb(resolved)
            return self.resolve_color(resolved, default)

        # Not found
        return self._hex_to_rgb(default)

    def resolve_color_hex(self, value: str | None, default: str = "#333333") -> str:
        """Resolve a color value to hex string (without #).

        Useful for XML operations that need raw hex values.

        Args:
            value: Color name or hex value
            default: Fallback hex color

        Returns:
            Hex string without # (e.g., "1E3A5F")
        """
        if not value:
            return default.lstrip("#")

        if value.startswith("#"):
            return value.lstrip("#")

        resolved = self.colors.get(value, default)
        if resolved.startswith("#"):
            return resolved.lstrip("#")

        # Recursive resolution for aliases
        return self.resolve_color_hex(resolved, default)

    def get_element_style(self, element: str) -> dict[str, Any]:
        """Get the style configuration for a named element.

        Args:
            element: Element name (title, heading1, body, bullet, etc.)

        Returns:
            Style dict with font, alignment, spacing, etc.
        """
        return self.styles.get(element, self.styles.get("body", {}))

    def get_font_config(self, element: str) -> dict[str, Any]:
        """Get the font configuration for an element.

        Args:
            element: Element name

        Returns:
            Font config dict with name, size, bold, italic, color, etc.
        """
        style = self.get_element_style(element)
        return style.get("font", {})

    def get_font_size(self, element: str, default: float = 11) -> Pt:
        """Get font size for an element as Pt.

        Args:
            element: Element name
            default: Default size in points

        Returns:
            Pt object for python-docx
        """
        font = self.get_font_config(element)
        return Pt(font.get("size", default))

    def get_font_name(self, element: str, default: str = "Arial") -> str:
        """Get font name for an element.

        Args:
            element: Element name
            default: Default font name

        Returns:
            Font family name
        """
        font = self.get_font_config(element)
        return font.get("name", default)

    def get_spacing(self, element: str, which: str, default: float = 0) -> Pt:
        """Get spacing value for an element.

        Args:
            element: Element name
            which: "before" or "after"
            default: Default spacing in points

        Returns:
            Pt object for python-docx
        """
        style = self.get_element_style(element)
        spacing = style.get("spacing", {})
        return Pt(spacing.get(which, default))

    def get_alignment(self, element: str) -> str | None:
        """Get alignment for an element.

        Args:
            element: Element name

        Returns:
            Alignment string ("left", "center", "right", "justify") or None
        """
        style = self.get_element_style(element)
        return style.get("alignment")

    def get_table_style(self, table_style: str = "default") -> dict[str, Any]:
        """Get table style configuration.

        Args:
            table_style: Table style name

        Returns:
            Table style dict with headerRow, bodyRow, etc.
        """
        return self.tables.get(table_style, self.tables.get("default", {}))

    def get_list_config(self, list_type: str = "bullet") -> dict[str, Any]:
        """Get list configuration.

        Args:
            list_type: List type (bullet, check, cross, dash)

        Returns:
            List config dict with symbol and color
        """
        return self.lists.get(list_type, self.lists.get("bullet", {}))

    def get_list_symbol(self, list_type: str = "bullet") -> str:
        """Get the symbol for a list type.

        Args:
            list_type: List type name

        Returns:
            Symbol string (e.g., "•", "✓")
        """
        config = self.get_list_config(list_type)
        return config.get("symbol", "\u2022")  # Default: bullet

    def get_page_margins(self) -> dict[str, Inches]:
        """Get page margins as Inches objects.

        Returns:
            Dict with top, right, bottom, left as Inches
        """
        margins = self.page.get("margins", {})
        result = {}

        for side in ("top", "right", "bottom", "left"):
            value = margins.get(side, "1in")
            # Parse "1in" or "0.75in" format
            if isinstance(value, str) and value.endswith("in"):
                result[side] = Inches(float(value.replace("in", "")))
            elif isinstance(value, (int, float)):
                result[side] = Inches(value)
            else:
                result[side] = Inches(1.0)

        return result

    def get_contact(self, key: str, default: str = "") -> str:
        """Get a contact info value.

        Args:
            key: Contact key (company, tagline, email, etc.)
            default: Default value

        Returns:
            Contact value string
        """
        return self.contact.get(key, default)

    def get_body_font(self) -> dict[str, Any]:
        """Get the default body font configuration.

        Returns:
            Font config dict from fonts.body
        """
        return self.fonts.get("body", {"name": "Arial", "size": 11})

    def get_heading_font(self) -> dict[str, Any]:
        """Get the default heading font configuration.

        Returns:
            Font config dict from fonts.heading
        """
        return self.fonts.get("heading", {"name": "Arial", "bold": True})

    def get_mono_font(self) -> dict[str, Any]:
        """Get the monospace font configuration.

        Returns:
            Font config dict from fonts.mono
        """
        return self.fonts.get("mono", {"name": "Courier New", "size": 10})

    @staticmethod
    def _hex_to_rgb(hex_color: str) -> RGBColor:
        """Convert hex color string to RGBColor.

        Args:
            hex_color: Hex string like "#1E3A5F" or "#333"

        Returns:
            RGBColor object
        """
        hex_color = hex_color.lstrip("#")

        # Handle shorthand (#333 -> #333333)
        if len(hex_color) == 3:
            hex_color = "".join(c * 2 for c in hex_color)

        return RGBColor(
            int(hex_color[0:2], 16),
            int(hex_color[2:4], 16),
            int(hex_color[4:6], 16)
        )

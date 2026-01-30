#!/usr/bin/env python3
"""High-level DOCX operations.

Provides convenient commands for common DOCX tasks:
- extract: Extract text from DOCX
- accept-changes: Accept all tracked changes
- to-pdf: Convert DOCX to PDF

Usage:
    ./run tool/docx_tools.py extract document.docx
    ./run tool/docx_tools.py accept-changes input.docx output.docx
    ./run tool/docx_tools.py to-pdf document.docx
"""

import argparse
import json
import sys
import tempfile
from pathlib import Path

# Add parent to path for imports
sys.path.insert(0, str(Path(__file__).parent))


def extract_text(docx_path: str, output_path: str = None,
                 format: str = 'plain', track_changes: str = 'accept') -> dict:
    """Extract text from a DOCX file using python-docx.

    Args:
        docx_path: Path to DOCX file
        output_path: Optional output file path
        format: Output format (plain, markdown, html)
        track_changes: How to handle tracked changes (accept, reject, all)

    Returns:
        Dict with extracted text or file path
    """
    from docx import Document

    docx_path = Path(docx_path)
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    # For formats other than plain, use mammoth for better conversion
    if format in ('markdown', 'md', 'html'):
        import mammoth

        with open(docx_path, 'rb') as f:
            if format == 'html':
                result = mammoth.convert_to_html(f)
            else:
                result = mammoth.convert_to_markdown(f)
            text = result.value
    else:
        # Plain text extraction with python-docx
        doc = Document(str(docx_path))

        paragraphs = []
        for para in doc.paragraphs:
            paragraphs.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                paragraphs.append('\t'.join(row_text))

        text = '\n'.join(paragraphs)

    # Write to file if output path specified
    if output_path:
        Path(output_path).write_text(text)
        return {
            'status': 'success',
            'operation': 'extract',
            'input': str(docx_path),
            'output': output_path,
            'format': format,
        }
    else:
        return {
            'status': 'success',
            'operation': 'extract',
            'input': str(docx_path),
            'format': format,
            'text': text,
        }


def accept_changes(input_path: str, output_path: str) -> dict:
    """Accept all tracked changes in a DOCX file.

    This unpacks the DOCX, runs accept_changes, and repacks.

    Args:
        input_path: Path to input DOCX
        output_path: Path for output DOCX

    Returns:
        Dict with operation results
    """
    from office.unpack import unpack_docx
    from office.accept_changes import accept_changes as do_accept
    from office.pack import pack_docx

    input_path = Path(input_path)
    output_path = Path(output_path)

    if not input_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {input_path}")

    # Create temp directory for unpacking
    with tempfile.TemporaryDirectory() as temp_dir:
        unpacked_dir = Path(temp_dir) / 'unpacked'

        # Unpack
        unpack_result = unpack_docx(str(input_path), str(unpacked_dir), merge_runs=False)

        # Accept changes
        accept_result = do_accept(str(unpacked_dir))

        # Repack
        pack_result = pack_docx(str(unpacked_dir), str(output_path))

    return {
        'status': 'success',
        'operation': 'accept_changes',
        'input': str(input_path),
        'output': str(output_path),
        'changes_accepted': accept_result.get('total_changes', 0),
    }


def convert_to_pdf(docx_path: str, output_path: str = None) -> dict:
    """Convert DOCX to PDF using mammoth + weasyprint.

    Args:
        docx_path: Path to DOCX file
        output_path: Optional output PDF path (defaults to same name with .pdf)

    Returns:
        Dict with operation results
    """
    import mammoth
    from weasyprint import HTML, CSS

    docx_path = Path(docx_path)
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    # Default output path
    if output_path is None:
        output_path = docx_path.with_suffix('.pdf')
    else:
        output_path = Path(output_path)

    # Convert DOCX to HTML with mammoth
    with open(docx_path, 'rb') as f:
        result = mammoth.convert_to_html(f)
        html_content = result.value

    # Wrap in basic HTML structure with styling
    full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
            font-size: 11pt;
            line-height: 1.5;
            max-width: 7in;
            margin: 0.75in auto;
            color: #333;
        }}
        h1 {{ font-size: 18pt; margin-top: 24pt; margin-bottom: 12pt; }}
        h2 {{ font-size: 14pt; margin-top: 18pt; margin-bottom: 10pt; }}
        h3 {{ font-size: 12pt; margin-top: 14pt; margin-bottom: 8pt; }}
        p {{ margin: 0 0 10pt 0; }}
        table {{ border-collapse: collapse; width: 100%; margin: 12pt 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8pt; text-align: left; }}
        th {{ background-color: #f5f5f5; font-weight: bold; }}
        ul, ol {{ margin: 0 0 10pt 0; padding-left: 20pt; }}
        li {{ margin-bottom: 4pt; }}
        img {{ max-width: 100%; height: auto; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""

    # Convert to PDF with WeasyPrint
    HTML(string=full_html).write_pdf(str(output_path))

    return {
        'status': 'success',
        'operation': 'to_pdf',
        'input': str(docx_path),
        'output': str(output_path),
    }


def main():
    parser = argparse.ArgumentParser(
        description='High-level DOCX operations',
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )

    subparsers = parser.add_subparsers(dest='command', help='Command to run')

    # extract command
    extract_parser = subparsers.add_parser('extract', help='Extract text from DOCX')
    extract_parser.add_argument('docx', help='Path to DOCX file')
    extract_parser.add_argument('-o', '--output', help='Output file path')
    extract_parser.add_argument('-f', '--format', default='plain',
                                choices=['plain', 'markdown', 'md', 'html'],
                                help='Output format (default: plain)')
    extract_parser.add_argument('--track-changes', default='accept',
                                choices=['accept', 'reject', 'all'],
                                help='Track changes handling (default: accept)')

    # accept-changes command
    accept_parser = subparsers.add_parser('accept-changes',
                                          help='Accept all tracked changes')
    accept_parser.add_argument('input', help='Input DOCX file')
    accept_parser.add_argument('output', help='Output DOCX file')

    # to-pdf command
    pdf_parser = subparsers.add_parser('to-pdf', help='Convert DOCX to PDF')
    pdf_parser.add_argument('docx', help='Path to DOCX file')
    pdf_parser.add_argument('-o', '--output', help='Output PDF path')

    # Global options
    parser.add_argument('--output-format', default='json',
                        choices=['json', 'text'],
                        help='Output format (default: json)')

    args = parser.parse_args()

    if not args.command:
        parser.print_help()
        sys.exit(1)

    try:
        if args.command == 'extract':
            result = extract_text(args.docx, args.output, args.format, args.track_changes)
        elif args.command == 'accept-changes':
            result = accept_changes(args.input, args.output)
        elif args.command == 'to-pdf':
            result = convert_to_pdf(args.docx, args.output)
        else:
            parser.print_help()
            sys.exit(1)

        if args.output_format == 'json':
            print(json.dumps(result, indent=2))
        else:
            if 'text' in result:
                print(result['text'])
            else:
                print(f"Operation: {result['operation']}")
                print(f"Input: {result['input']}")
                if 'output' in result:
                    print(f"Output: {result['output']}")

    except Exception as e:
        error_result = {
            'status': 'error',
            'operation': args.command,
            'error': str(e)
        }
        print(json.dumps(error_result, indent=2), file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()
